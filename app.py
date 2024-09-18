from flask import Flask, request, jsonify, render_template
import requests

import google.generativeai as genai
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
from dotenv import load_dotenv

from dotenv import load_dotenv
import os
app = Flask(__name__)

resume = """
Rangel Anselm Koli
+1 315-374-9529 | rangelkoli@gmail.com | linkedin.com/in/rangelkoli | github.com/rangelkoli Career Objective
I aim to secure a dynamic role in a forward-thinking organization where I can leverage my technical expertise to drive innovation, contribute to cutting-edge technologies, and advance both my professional skills and the company’s objectives.
    Technical Skills
Languages: Python, Java, Javascript, C++, C#
Frameworks: React Native, React, .NET, Django, FastAPI, Flask, NodeJS Libraries: Tensorflow, Keras, OpenCV, Pytorch
Education
Syracuse University
Master of Science in Computer Science
Experience
App Developer Intern
   Delet
•
•
Syracuse, NY
Anticipated May 2025
June 2024 - August 2024
San Francisco, CA
Transformed a company website into a high-performing React Native app, ensuring seamless functionality and user experience across both platforms by adhering to industry best practices and optimizing API integrations. Collaborated with cross-functional teams to design, develop, and deploy the React Native app, resulting in a 20% increase in user engagement and retention by implementing efficient API calls and maintaining high code quality standards.
Intern March 2022 - May 2022 GRIP (The Sparks Foundation) Mumbai, MH
• Spearheaded the design and development of a user-friendly charity website, integrating a secure payment gateway with React, resulting in a 9% increase in online donations by optimizing user flow and ensuring transaction security.
• Partnered with cross-functional teams to revamp website design and functionality, achieving a 23% increase in user engagement by improving UI/UX and effectively communicating the charity’s mission.
Projects
WalkSafe | React-Native, Flask, Supabase, Google Cloud June 2020 – Present
• Developed and integrated a real-time tracking feature using React-Native, boosting user connectivity and peace of mind by enabling effortless monitoring of friends and family on an interactive map, which led to a 15% increase
in app engagement.
• Engineered a route optimization algorithm with Flask, enhancing user safety by analyzing crime data to suggest
secure paths, resulting in a 20% decrease in user-reported safety incidents.
• Designed and implemented a Check-In Service using Supabase, automating notifications upon reaching designated
destinations, which increased user satisfaction by 18% and enhanced the sense of security among users.
Anonimo | React-Native, Flask, Supabase, Google Cloud June 2020 – Present
• Applied Machine Learning techniques to develop a user-friendly social platform, fostering interaction and pro-
viding support for users in distress, leading to a 20% increase in user retention and engagement.
• Deployed a robust Machine Learning model, classifying users into categories of neutrality or desolation, enhancing
the platform’s ability to offer targeted support and increasing user satisfaction by 15%.
Extra-Curricular
• Secured 1st place in the Code Crush 1.0 Hackathon at St. Francis Institute of Technology by developing an innovative machine learning solution, outperforming 30+ competing teams and demonstrating superior problem- solving skills.
• Achieved 1st place in the CuseHacks Beta 2024 Hackathon at Syracuse University by creating a cutting-edge AI- based application, surpassing 50+ teams and showcasing advanced technical skills.
• Produced and managed a YouTube channel, creating educational and technical content that attracted and engaged 15K active subscribers, driving a 20% monthly growth rate in viewership.
  
"""


def generateAttentionHook(job_description):
    genai.configure(api_key=os.environ.get("API_KEY"))

    model = genai.GenerativeModel("gemini-1.5-flash")

    chat = model.start_chat(
        history=[
        ]
    )

    initial_prompt = "Based on this job description, what is the biggest challenge someone in this position would face day-to-day? \n"
    response = chat.send_message(initial_prompt + job_description)

    return response.text




@app.route("/coverLetter", methods=["POST", "GET"])
def generate_cover_letter():
    if request.method == "POST":
        genai.configure(api_key=os.environ.get("API_KEY"))

        job_description = request.form.get("job-description")
        job_title = request.form.get("job-title")

        job_link = request.form.get("job-link")
        print(job_link)
        # Linekdin URL
        response = requests.get(job_link)

        if response.status_code == 200:
            soup = BeautifulSoup(response.text, 'html.parser')
            company_name = soup.find("h1", {"class": "top-card-layout__title"}).text


            job_title = soup.find("a", {"class": "topcard__org-name-link"}).text

            job_description = soup.find("div", {"class": "show-more-less-html__markup"}).get_text()

        model = genai.GenerativeModel("gemini-1.5-flash")
        chat = model.start_chat(
            history=[
                    {"role": "user", "parts": "Hello"},
                    {"role": "model", "parts": generateAttentionHook(job_description)},

            ]
        )
        # initial_prompt = "Based on this job description, what is the biggest challenge someone in this position would face day-to-day? \n"
        # response = chat.send_message(initial_prompt + job_description)
        # print(chat.history)
        response = chat.send_message("""
    Based on the above job description,  the challenges you provided above and the """ + resume + """
    Write an attention-grabbing hook for your cover letter that highlights your experiences and qualifications in a way that shows you empathize and can successfully take on the challenges of a """ + job_title + """ role.
    Consider incorporating specific examples of how you've tackled these challenges in your past work and explore creative ways to express your enthusiasm for the opportunity. Keep your hook within 100 words and only have 1 paragraph
                                    """, )
        attentionHook = response.text

        response = chat.send_message("""
    Write the body of your cover letter that elaborates on your experiences and qualifications, aligning them with the job requirements and company culture.
                                     Here are some details you might need Job title : """ + job_title + """ Company Name : """ + company_name + """ 
                                     This is what you have so far: """ + attentionHook + """
                                     DO NOT WRITE THE ATTENTION HOOK AGAIN, but give me the other paragraphs 
    Highlight your achievements and skills, and explain how they make you a perfect fit for the """ + job_title + """ role.
    Use all the experiences and qualifications in the resume above to write the body of the cover letter. DO NOT HAVE ANY BULLET POINTS, or variables so that the cover letter is complete.
    Include specific examples of your work that demonstrate your expertise and passion for the field. Make sure the cover letter is less than 500 words and will fit on one page.
                                    """, )
    
        document = Document(docx=None)

        document.add_heading('Cover Letter', 1)
        document.add_heading('Rangel Anselm Koli', level=1)
        p = document.add_paragraph(attentionHook, style='BodyText')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run("\n")
        s = document.add_paragraph(response.text, style='BodyText')
        s.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        docName = f"{job_title}CoverLetter.docx"
        document.save(
            docName
        )
        return jsonify({"cover_letter": response.text, "attentionHook": attentionHook, "company_name": company_name, "job_title": job_title})
    return render_template("index.html")

if __name__ == "__main__":
    app.run(debug=True)
